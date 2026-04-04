"""SaintSal Labs — Career Suite v2 (Supabase-Backed)
All persistence backed by Supabase. PDF/DOCX export. SAL coaching. DNA autofill.
Column names match the actual Supabase schema exactly.
"""
from fastapi import APIRouter, Request, UploadFile, File
from fastapi.responses import JSONResponse, StreamingResponse
import os, json, uuid, httpx, io, shutil
from datetime import datetime, timezone
from pathlib import Path

router = APIRouter(prefix="/api/career/v2", tags=["career-v2"])

def _sb():
    from supabase import create_client
    url = os.environ.get("SUPABASE_URL", "")
    key = os.environ.get("SUPABASE_SERVICE_KEY", "")
    if not url or not key:
        return None
    return create_client(url, key)

def _user_id(request: Request):
    if hasattr(request.state, 'user_id'):
        return str(request.state.user_id)
    auth = request.headers.get("authorization", "")
    if auth.startswith("Bearer "):
        return auth.split(" ", 1)[1][:100]
    return request.query_params.get("user_id", "anonymous")

def _now():
    return datetime.now(timezone.utc).isoformat()

def _clean(rows):
    for r in rows:
        r.pop("id", None)
    return rows

VALID_TONES = {"professional", "executive", "conversational", "creative"}
def _map_tone(style):
    if style in VALID_TONES:
        return style
    return "professional"


# ── 1. CAREER PROFILE ──

@router.get("/profile")
async def get_career_profile(request: Request):
    uid = _user_id(request)
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    try:
        r = sb.table("career_profiles").select("*").eq("user_id", uid).limit(1).execute()
        if r.data:
            row = r.data[0]
            row.pop("id", None)
            return row
        profile = {"user_id": uid}
        sb.table("career_profiles").insert(profile).execute()
        return {"user_id": uid, "full_name": None, "onboarding_completed": False}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.post("/profile")
async def save_career_profile(request: Request):
    uid = _user_id(request)
    body = await request.json()
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    VALID = {"full_name","email","phone","location","linkedin_url","portfolio_url","website_url",
             "github_url","headline","summary","years_experience","current_title","current_company",
             "industry","target_role","target_salary_min","target_salary_max","target_locations",
             "remote_preference","job_search_status","skills","certifications","education","languages",
             "headshot_url","background_image_url","custom_logo_url","business_dna_id"}
    payload = {k: v for k, v in body.items() if k in VALID and v is not None}
    payload["user_id"] = uid
    payload["updated_at"] = _now()
    try:
        sb.table("career_profiles").upsert(payload, on_conflict="user_id").execute()
        return {"success": True}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.post("/profile/autofill-dna")
async def autofill_from_dna(request: Request):
    uid = _user_id(request)
    body = await request.json()
    dna = body.get("dna", {})
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    profile = {
        "user_id": uid,
        "full_name": ((dna.get("first_name","") + " " + dna.get("last_name","")).strip()) or None,
        "email": dna.get("email") or None,
        "phone": dna.get("phone") or None,
        "headline": dna.get("tagline") or None,
        "summary": dna.get("bio") or None,
        "current_company": dna.get("business_name") or None,
        "industry": dna.get("industry") or None,
        "location": ((dna.get("business_city","") + ", " + dna.get("business_state","")).strip(", ")) or None,
        "years_experience": dna.get("years_in_business"),
        "website_url": dna.get("website") or None,
        "updated_at": _now(),
    }
    profile = {k: v for k, v in profile.items() if v is not None}
    try:
        sb.table("career_profiles").upsert(profile, on_conflict="user_id").execute()
        return {"success": True, "autofilled": list(profile.keys())}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ── 2. RESUMES ──

@router.get("/resumes")
async def list_resumes(request: Request):
    uid = _user_id(request)
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    try:
        r = sb.table("resumes").select("*").eq("user_id", uid).order("updated_at", desc=True).execute()
        data = r.data or []
        for row in data:
            row["resume_id"] = row.pop("id", "")
        return {"resumes": data, "total": len(data)}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.post("/resumes")
async def create_resume(request: Request):
    uid = _user_id(request)
    body = await request.json()
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    resume_id = str(uuid.uuid4())
    # Map frontend fields to actual schema columns
    contact = {
        "full_name": body.get("full_name", ""),
        "email": body.get("email", ""),
        "phone": body.get("phone", ""),
        "location": body.get("location", ""),
        "linkedin": body.get("linkedin_url", ""),
        "website": body.get("website", ""),
    }
    resume = {
        "id": resume_id,
        "user_id": uid,
        "name": body.get("resume_name", body.get("name", "Untitled Resume")),
        "template_id": body.get("template_style", "professional"),
        "is_primary": body.get("is_primary", False),
        "contact_info": json.dumps(contact),
        "summary_text": body.get("summary", ""),
        "skills_section": json.dumps(body.get("skills", [])),
        "education_section": json.dumps(body.get("education", [])),
        "certifications_section": json.dumps(body.get("certifications", [])),
        "raw_content": json.dumps(body.get("work_experience", [])),
        "target_job_description": body.get("target_role", ""),
        "created_at": _now(),
        "updated_at": _now(),
    }
    try:
        sb.table("resumes").insert(resume).execute()
        # Save work history entries
        for idx, job in enumerate(body.get("work_experience", [])):
            wh = {
                "id": str(uuid.uuid4()),
                "resume_id": resume_id,
                "user_id": uid,
                "company_name": job.get("company", ""),
                "job_title": job.get("title", ""),
                "start_date": job.get("start_date"),
                "end_date": job.get("end_date"),
                "is_current": job.get("is_current", False),
                "description": job.get("description", ""),
                "achievements": json.dumps(job.get("achievements", [])),
                "sort_order": idx,
            }
            wh = {k: v for k, v in wh.items() if v is not None}
            try:
                sb.table("resume_work_history").insert(wh).execute()
            except Exception:
                pass
        return {"success": True, "resume_id": resume_id}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.put("/resumes/{resume_id}")
async def update_resume(resume_id: str, request: Request):
    uid = _user_id(request)
    body = await request.json()
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    VALID = {"name","template_id","is_primary","contact_info","summary_text","skills_section",
             "education_section","certifications_section","raw_content","enhanced_content",
             "target_job_description","background_image_url","header_style","font_family","color_scheme"}
    update = {k: v for k, v in body.items() if k in VALID}
    update["updated_at"] = _now()
    try:
        sb.table("resumes").update(update).eq("id", resume_id).eq("user_id", uid).execute()
        return {"success": True}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.post("/resumes/{resume_id}/enhance")
async def enhance_resume_with_sal(resume_id: str, request: Request):
    uid = _user_id(request)
    body = await request.json()
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    try:
        r = sb.table("resumes").select("*").eq("id", resume_id).eq("user_id", uid).limit(1).execute()
        if not r.data:
            return JSONResponse({"error": "Resume not found"}, status_code=404)
        resume = r.data[0]
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    # Also fetch work history
    try:
        wh = sb.table("resume_work_history").select("*").eq("resume_id", resume_id).order("sort_order").execute()
        work = wh.data or []
    except Exception:
        work = []

    contact = json.loads(resume.get("contact_info", "{}")) if isinstance(resume.get("contact_info"), str) else (resume.get("contact_info") or {})
    skills = json.loads(resume.get("skills_section", "[]")) if isinstance(resume.get("skills_section"), str) else (resume.get("skills_section") or [])
    education = json.loads(resume.get("education_section", "[]")) if isinstance(resume.get("education_section"), str) else (resume.get("education_section") or [])

    resume_text = f"""Name: {contact.get('full_name','')}
Summary: {resume.get('summary_text','')}
Skills: {', '.join(skills) if isinstance(skills, list) else str(skills)}
Education: {json.dumps(education)}
Work Experience: {json.dumps([{"company": w.get("company_name",""), "title": w.get("job_title",""), "description": w.get("description",""), "achievements": json.loads(w.get("achievements","[]")) if isinstance(w.get("achievements"), str) else w.get("achievements",[])} for w in work])}"""

    target_role = body.get("target_role", resume.get("target_job_description", ""))
    ANTHROPIC_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
    if not ANTHROPIC_KEY:
        return JSONResponse({"error": "AI not configured"}, status_code=500)

    system = """You are SaintSal, an elite career AI. Enhance this resume to be ATS-optimized, achievement-focused, and compelling.
Return ONLY valid JSON with: enhanced_summary (string), enhanced_skills (array of strings),
enhanced_experience (array of {company, title, description, achievements: array}),
ats_score (0-100), keywords_added (array), improvement_notes (string)."""

    prompt = f"Enhance this resume{' for: ' + target_role if target_role else ''}:\n\n{resume_text}\n\nMake it powerful. Quantify achievements. Add industry keywords."

    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post("https://api.anthropic.com/v1/messages",
                headers={"x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"},
                json={"model": "claude-sonnet-4-20250514", "max_tokens": 4096,
                      "system": system, "messages": [{"role": "user", "content": prompt}]})
            ai_text = resp.json().get("content", [{}])[0].get("text", "{}")
            enhanced = json.loads(ai_text.strip().removeprefix("```json").removesuffix("```").strip())
    except Exception as e:
        return JSONResponse({"error": f"AI failed: {str(e)}"}, status_code=500)

    try:
        sb.table("resumes").update({
            "enhanced_content": json.dumps(enhanced),
            "saintssal_enhanced": True,
            "ats_score": enhanced.get("ats_score", 0),
            "ats_keywords": json.dumps(enhanced.get("keywords_added", [])),
            "enhancement_prompt": prompt[:500],
            "updated_at": _now(),
        }).eq("id", resume_id).execute()
    except Exception:
        pass

    return {"success": True, "enhanced": enhanced, "resume_id": resume_id}


@router.get("/resumes/{resume_id}/export/{fmt}")
async def export_resume(resume_id: str, fmt: str, request: Request):
    uid = _user_id(request)
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    try:
        r = sb.table("resumes").select("*").eq("id", resume_id).eq("user_id", uid).limit(1).execute()
        if not r.data:
            return JSONResponse({"error": "Resume not found"}, status_code=404)
        resume = r.data[0]
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    # Also fetch work history
    try:
        wh = sb.table("resume_work_history").select("*").eq("resume_id", resume_id).order("sort_order").execute()
        work = wh.data or []
    except Exception:
        work = []

    contact = json.loads(resume.get("contact_info","{}")) if isinstance(resume.get("contact_info"), str) else (resume.get("contact_info") or {})
    skills = json.loads(resume.get("skills_section","[]")) if isinstance(resume.get("skills_section"), str) else (resume.get("skills_section") or [])
    education = json.loads(resume.get("education_section","[]")) if isinstance(resume.get("education_section"), str) else (resume.get("education_section") or [])

    # Use enhanced content if available
    enhanced = json.loads(resume.get("enhanced_content","{}")) if isinstance(resume.get("enhanced_content"), str) else (resume.get("enhanced_content") or {})
    name = contact.get("full_name", "Resume")
    headline = enhanced.get("enhanced_headline", "")
    summary = enhanced.get("enhanced_summary") or resume.get("summary_text", "")
    final_skills = enhanced.get("enhanced_skills") or skills
    final_experience = enhanced.get("enhanced_experience") or [
        {"company": w.get("company_name",""), "title": w.get("job_title",""),
         "description": w.get("description",""),
         "achievements": json.loads(w.get("achievements","[]")) if isinstance(w.get("achievements"), str) else w.get("achievements",[])}
        for w in work
    ]
    email = contact.get("email", "")
    phone = contact.get("phone", "")
    location = contact.get("location", "")
    filename = f"{name.replace(' ', '_')}_Resume"

    if fmt == "pdf":
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.colors import HexColor

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.6*inch, bottomMargin=0.5*inch,
                                leftMargin=0.7*inch, rightMargin=0.7*inch)
        styles = getSampleStyleSheet()
        gold = HexColor("#D4AF37")
        dark = HexColor("#1a1a1a")
        name_s = ParagraphStyle("N", parent=styles["Title"], fontSize=22, textColor=dark, spaceAfter=4, fontName="Helvetica-Bold")
        hl_s = ParagraphStyle("HL", parent=styles["Normal"], fontSize=11, textColor=HexColor("#555"), spaceAfter=6)
        ct_s = ParagraphStyle("CT", parent=styles["Normal"], fontSize=9, textColor=HexColor("#666"), spaceAfter=12)
        sec_s = ParagraphStyle("SC", parent=styles["Heading2"], fontSize=12, textColor=gold, fontName="Helvetica-Bold", spaceAfter=6, spaceBefore=14)
        bd_s = ParagraphStyle("BD", parent=styles["Normal"], fontSize=10, textColor=dark, spaceAfter=4, leading=14)
        bl_s = ParagraphStyle("BL", parent=bd_s, leftIndent=12, bulletIndent=0, spaceAfter=2)

        story = [Paragraph(name, name_s)]
        if headline:
            story.append(Paragraph(headline, hl_s))
        cp = [p for p in [email, phone, location] if p]
        if cp:
            story.append(Paragraph(" | ".join(cp), ct_s))
        story.append(HRFlowable(width="100%", thickness=1, color=gold, spaceAfter=10))
        if summary:
            story.append(Paragraph("PROFESSIONAL SUMMARY", sec_s))
            story.append(Paragraph(summary, bd_s))
        if final_experience:
            story.append(Paragraph("EXPERIENCE", sec_s))
            for job in final_experience:
                t = job.get("title", job.get("job_title", ""))
                c = job.get("company", job.get("company_name", ""))
                story.append(Paragraph(f"<b>{t}</b> — {c}", bd_s))
                if job.get("description"):
                    story.append(Paragraph(job["description"], bd_s))
                for a in job.get("achievements", []):
                    if isinstance(a, str):
                        story.append(Paragraph(f"• {a}", bl_s))
                story.append(Spacer(1, 4))
        if education:
            story.append(Paragraph("EDUCATION", sec_s))
            for e in education:
                if isinstance(e, dict):
                    story.append(Paragraph(f"<b>{e.get('degree','')}</b> — {e.get('school','')}", bd_s))
                else:
                    story.append(Paragraph(str(e), bd_s))
        if final_skills:
            story.append(Paragraph("SKILLS", sec_s))
            story.append(Paragraph(", ".join(final_skills) if isinstance(final_skills, list) else str(final_skills), bd_s))
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#ccc"), spaceAfter=6))
        sal_s = ParagraphStyle("SAL", parent=styles["Normal"], fontSize=7, textColor=HexColor("#999"), alignment=1)
        story.append(Paragraph("Enhanced by SaintSal Labs | saintsallabs.com", sal_s))
        doc.build(story)
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/pdf",
                                 headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'})

    elif fmt == "docx":
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.styles["Normal"].font.size = Pt(10)
        p = doc.add_heading(name, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if headline:
            doc.add_paragraph(headline)
        cp = [x for x in [email, phone, location] if x]
        if cp:
            doc.add_paragraph(" | ".join(cp))
        if summary:
            doc.add_heading("PROFESSIONAL SUMMARY", level=2)
            doc.add_paragraph(summary)
        if final_experience:
            doc.add_heading("EXPERIENCE", level=2)
            for job in final_experience:
                t = job.get("title", job.get("job_title", ""))
                c = job.get("company", job.get("company_name", ""))
                p = doc.add_paragraph()
                run = p.add_run(f"{t} — {c}")
                run.bold = True
                if job.get("description"):
                    doc.add_paragraph(job["description"])
                for a in job.get("achievements", []):
                    if isinstance(a, str):
                        doc.add_paragraph(f"• {a}")
        if education:
            doc.add_heading("EDUCATION", level=2)
            for e in education:
                if isinstance(e, dict):
                    doc.add_paragraph(f"{e.get('degree','')} — {e.get('school','')}")
                else:
                    doc.add_paragraph(str(e))
        if final_skills:
            doc.add_heading("SKILLS", level=2)
            doc.add_paragraph(", ".join(final_skills) if isinstance(final_skills, list) else str(final_skills))
        p = doc.add_paragraph("\nEnhanced by SaintSal Labs | saintsallabs.com")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'})

    return JSONResponse({"error": "Format must be 'pdf' or 'docx'"}, status_code=400)


# ── 3. JOB TRACKER (Supabase) ──

@router.get("/tracker")
async def get_job_tracker(request: Request):
    uid = _user_id(request)
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    try:
        r = sb.table("job_applications").select("*").eq("user_id", uid).order("updated_at", desc=True).execute()
        jobs = r.data or []
        for j in jobs:
            j["job_id"] = j.pop("id", "")
        kanban = {}
        # Valid Supabase statuses: saved, applied, phone_screen, rejected
        for status in ["saved","applied","phone_screen","rejected"]:
            kanban[status] = [j for j in jobs if j.get("status") == status]
        return {"kanban": kanban, "total": len(jobs), "jobs": jobs}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.post("/tracker")
async def add_to_tracker(request: Request):
    uid = _user_id(request)
    body = await request.json()
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    app_id = str(uuid.uuid4())
    # Map to actual schema columns
    job = {
        "id": app_id,
        "user_id": uid,
        "job_title": body.get("job_title", body.get("title", "")),
        "company_name": body.get("company_name", body.get("company", "")),
        "job_url": body.get("job_url", body.get("url", "")),
        "location": body.get("location", ""),
        "salary_range": body.get("salary_range", ""),
        "remote_type": body.get("remote_type", "onsite"),
        "status": body.get("status", "saved"),  # Valid: saved, applied, phone_screen, rejected
        "notes": body.get("notes", ""),
        "job_source": body.get("source", body.get("job_source", "manual")),
        "applied_date": body.get("applied_date"),
        "job_description": body.get("description", body.get("job_description", "")),
        "recruiter_name": body.get("recruiter_name", ""),
        "recruiter_email": body.get("recruiter_email", ""),
        "created_at": _now(),
        "updated_at": _now(),
    }
    job = {k: v for k, v in job.items() if v is not None}
    try:
        sb.table("job_applications").insert(job).execute()
        return {"success": True, "job_id": app_id}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.put("/tracker/{job_id}")
async def update_tracker_job(job_id: str, request: Request):
    uid = _user_id(request)
    body = await request.json()
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    VALID = {"job_title","company_name","job_url","location","salary_range","remote_type","status",
             "notes","job_source","applied_date","response_date","recruiter_name","recruiter_email",
             "recruiter_phone","hiring_manager_name","hiring_manager_email","resume_id",
             "cover_letter_id","follow_up_date","follow_up_notes","offer_amount","offer_details",
             "rejection_reason","job_description"}
    update = {k: v for k, v in body.items() if k in VALID}
    update["updated_at"] = _now()
    if update.get("status") == "applied" and "applied_date" not in update:
        update["applied_date"] = _now()
    try:
        sb.table("job_applications").update(update).eq("id", job_id).eq("user_id", uid).execute()
        return {"success": True}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.delete("/tracker/{job_id}")
async def delete_tracker_job(job_id: str, request: Request):
    uid = _user_id(request)
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    try:
        sb.table("job_applications").delete().eq("id", job_id).eq("user_id", uid).execute()
        return {"success": True}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ── 4. INTERVIEWS ──

@router.get("/interviews/{job_id}")
async def get_interviews(job_id: str, request: Request):
    uid = _user_id(request)
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    try:
        r = sb.table("interviews").select("*").eq("job_application_id", job_id).eq("user_id", uid).order("interview_date").execute()
        data = r.data or []
        for row in data:
            row["interview_id"] = row.pop("id", "")
        return {"interviews": data, "total": len(data)}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.post("/interviews")
async def add_interview(request: Request):
    uid = _user_id(request)
    body = await request.json()
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    iid = str(uuid.uuid4())
    interview = {
        "id": iid,
        "user_id": uid,
        "job_application_id": body.get("job_id", body.get("job_application_id", "")),
        "interview_type": body.get("interview_type", "phone_screen"),
        "interview_date": body.get("interview_date"),
        "interviewer_name": body.get("interviewer_name", ""),
        "interviewer_title": body.get("interviewer_title", ""),
        "notes": body.get("notes", ""),
        "status": body.get("status", "scheduled"),
        "created_at": _now(),
    }
    interview = {k: v for k, v in interview.items() if v is not None}
    try:
        sb.table("interviews").insert(interview).execute()
        try:
            sb.table("job_applications").update({"status": "interview", "updated_at": _now()}).eq("id", interview.get("job_application_id","")).execute()
        except Exception:
            pass
        return {"success": True, "interview_id": iid}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.post("/interviews/{interview_id}/coach")
async def interview_coaching(interview_id: str, request: Request):
    uid = _user_id(request)
    body = await request.json()
    sb = _sb()
    company = body.get("company", "")
    role = body.get("role", "")
    stage = body.get("stage", "prep")
    interview_type = body.get("interview_type", "phone screen")

    # Try to get job details from DB
    if sb and body.get("job_id"):
        try:
            jr = sb.table("job_applications").select("company_name,job_title").eq("id", body["job_id"]).limit(1).execute()
            if jr.data:
                company = company or jr.data[0].get("company_name", "")
                role = role or jr.data[0].get("job_title", "")
        except Exception:
            pass

    ANTHROPIC_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
    if not ANTHROPIC_KEY:
        return JSONResponse({"error": "AI not configured"}, status_code=500)

    prompts = {
        "prep": f"Prepare me for a {interview_type} interview at {company} for {role}. Give: company research, 10 likely questions with STAR answers, questions to ask them, checklist, red flags. Return JSON: company_research, likely_questions (array of {{question, answer_framework}}), questions_to_ask, preparation_checklist, red_flags, confidence_tips.",
        "during": f"I'm in a {interview_type} at {company} for {role}. They asked: {body.get('question','')}. Help me answer using STAR. Return JSON: suggested_answer, key_points, follow_up_question_to_ask, body_language_tip.",
        "followup": f"My {interview_type} at {company} for {role} just ended. Write: thank-you email, LinkedIn message, follow-up plan. Return JSON: thank_you_email, linkedin_message, follow_up_timeline (array of {{day, action}}), next_steps.",
    }
    system = "You are SaintSal, an elite career coach. Give specific, actionable guidance. Return ONLY valid JSON."

    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post("https://api.anthropic.com/v1/messages",
                headers={"x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"},
                json={"model": "claude-sonnet-4-20250514", "max_tokens": 4096,
                      "system": system, "messages": [{"role": "user", "content": prompts.get(stage, prompts["prep"])}]})
            ai_text = resp.json().get("content", [{}])[0].get("text", "{}")
            coaching = json.loads(ai_text.strip().removeprefix("```json").removesuffix("```").strip())
    except json.JSONDecodeError:
        coaching = {"raw_response": ai_text}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    if sb:
        try:
            sb.table("interviews").update({"sal_prep_notes": json.dumps(coaching)}).eq("id", interview_id).execute()
        except Exception:
            pass

    return {"success": True, "stage": stage, "coaching": coaching}


# ── 5. EMAIL SIGNATURES ──

@router.get("/signatures")
async def list_signatures(request: Request):
    uid = _user_id(request)
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    try:
        r = sb.table("email_signatures").select("*").eq("user_id", uid).order("created_at", desc=True).execute()
        data = r.data or []
        for row in data:
            row["sig_id"] = row.pop("id", "")
        return {"signatures": data, "total": len(data)}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.post("/signatures")
async def save_signature(request: Request):
    uid = _user_id(request)
    body = await request.json()
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    sig = {
        "id": str(uuid.uuid4()),
        "user_id": uid,
        "name": body.get("signature_name", body.get("name", "My Signature")),
        "template_id": body.get("template_style", body.get("template_id", "professional")),
        "full_name": body.get("full_name", ""),
        "title": body.get("title", ""),
        "company": body.get("company", ""),
        "email": body.get("email", ""),
        "phone": body.get("phone", ""),
        "website": body.get("website", ""),
        "linkedin_url": body.get("linkedin_url", ""),
        "twitter_url": body.get("twitter_url", ""),
        "instagram_url": body.get("instagram_url", ""),
        "company_logo_url": body.get("logo_url", body.get("company_logo_url", "")),
        "banner_url": body.get("banner_url", ""),
        "custom_background_url": body.get("custom_background_url", ""),
        "color_primary": body.get("color_primary", "#D4AF37"),
        "color_accent": body.get("color_accent", body.get("color_secondary", "#1a1a1a")),
        "html_output": body.get("html_content", body.get("html_output", "")),
        "is_primary": body.get("is_primary", False),
        "created_at": _now(),
    }
    try:
        sb.table("email_signatures").insert(sig).execute()
        return {"success": True, "sig_id": sig["id"]}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ── 6. COVER LETTERS ──

@router.post("/cover-letters")
async def save_cover_letter(request: Request):
    uid = _user_id(request)
    body = await request.json()
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    cl = {
        "id": str(uuid.uuid4()),
        "user_id": uid,
        "job_application_id": body.get("job_id", body.get("job_application_id")),
        "resume_id": body.get("resume_id"),
        "name": body.get("name", "Cover Letter"),
        "body_raw": body.get("content", body.get("body_raw", "")),
        "body_enhanced": body.get("body_enhanced", ""),
        "tone": _map_tone(body.get("tone", body.get("style", "professional"))),
        "company_name": body.get("target_company", body.get("company_name", "")),
        "job_title": body.get("target_role", body.get("job_title", "")),
        "saintssal_enhanced": body.get("saintssal_enhanced", False),
        "created_at": _now(),
    }
    cl = {k: v for k, v in cl.items() if v is not None}
    try:
        sb.table("cover_letters").insert(cl).execute()
        return {"success": True, "cover_letter_id": cl["id"]}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.get("/cover-letters")
async def list_cover_letters(request: Request):
    uid = _user_id(request)
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    try:
        r = sb.table("cover_letters").select("*").eq("user_id", uid).order("created_at", desc=True).execute()
        data = r.data or []
        for row in data:
            row["cl_id"] = row.pop("id", "")
        return {"cover_letters": data, "total": len(data)}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ── 7. STAGE COACHING (Start → Job Won) ──

@router.post("/coach/stage-guidance")
async def stage_guidance(request: Request):
    body = await request.json()
    status = body.get("status", "wishlist")
    company = body.get("company", "")
    role = body.get("role", "")

    ANTHROPIC_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
    if not ANTHROPIC_KEY:
        return JSONResponse({"error": "AI not configured"}, status_code=500)

    guidance = {
        "wishlist": f"I'm interested in {role} at {company}. Pre-application checklist and company research.",
        "applied": f"Just applied for {role} at {company}. Follow-up email template, timing, and what to do while waiting.",
        "phone_screen": f"Phone screen for {role} at {company}. Key questions, answers, and phone interview tips.",
        "interview": f"In-person interview for {role} at {company}. STAR answers for 5 likely questions, body language, day-of checklist.",
        "offer": f"Offer for {role} at {company}. Negotiation script, benefits to negotiate, decision framework.",
        "accepted": f"Accepted {role} at {company}! 30-60-90 day plan, first-week tips, strong impression strategy.",
        "rejected": f"Rejected for {role} at {company}. Graceful response, lessons, keeping the door open.",
    }
    system = "You are SaintSal career coach. Be specific and encouraging. Return JSON: guidance, action_items (array), templates (array of {type, content}), timeline (array of {day, action}), motivation."

    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post("https://api.anthropic.com/v1/messages",
                headers={"x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"},
                json={"model": "claude-sonnet-4-20250514", "max_tokens": 4096,
                      "system": system, "messages": [{"role": "user", "content": guidance.get(status, guidance["wishlist"])}]})
            ai_text = resp.json().get("content", [{}])[0].get("text", "{}")
            coaching = json.loads(ai_text.strip().removeprefix("```json").removesuffix("```").strip())
    except json.JSONDecodeError:
        coaching = {"guidance": ai_text, "action_items": [], "motivation": "Keep pushing!"}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    return {"success": True, "status": status, "coaching": coaching}


# ── 8. SAVED JOB SEARCHES ──

@router.post("/saved-searches")
async def save_job_search(request: Request):
    uid = _user_id(request)
    body = await request.json()
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    search = {
        "id": str(uuid.uuid4()),
        "user_id": uid,
        "search_name": body.get("search_name", ""),
        "keywords": body.get("query", body.get("keywords", "")),
        "location": body.get("location", ""),
        "job_type": body.get("job_type", ""),
        "remote_only": body.get("remote_only", False),
        "salary_min": body.get("salary_min"),
        "salary_max": body.get("salary_max"),
        "alert_enabled": body.get("alert_enabled", False),
        "created_at": _now(),
    }
    search = {k: v for k, v in search.items() if v is not None}
    try:
        sb.table("job_saved_searches").insert(search).execute()
        return {"success": True}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.get("/saved-searches")
async def list_saved_searches(request: Request):
    uid = _user_id(request)
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    try:
        r = sb.table("job_saved_searches").select("*").eq("user_id", uid).order("created_at", desc=True).execute()
        data = r.data or []
        for row in data:
            row.pop("id", None)
        return {"saved_searches": data, "total": len(data)}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ── 9. COVER LETTER EXPORT (PDF / DOCX) ──

@router.get("/cover-letters/{cl_id}/export/{fmt}")
async def export_cover_letter(cl_id: str, fmt: str, request: Request):
    uid = _user_id(request)
    sb = _sb()
    if not sb:
        return JSONResponse({"error": "Supabase not configured"}, status_code=500)
    try:
        r = sb.table("cover_letters").select("*").eq("id", cl_id).eq("user_id", uid).limit(1).execute()
        if not r.data:
            return JSONResponse({"error": "Cover letter not found"}, status_code=404)
        cl = r.data[0]
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    body_text = cl.get("body_enhanced") or cl.get("body_raw", "")
    company = cl.get("company_name", "")
    role = cl.get("job_title", "")
    name = cl.get("name", "Cover Letter")
    filename = f"{name.replace(' ', '_')}_{company.replace(' ', '_')}" if company else name.replace(' ', '_')

    if fmt == "pdf":
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.colors import HexColor

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.8*inch, bottomMargin=0.6*inch,
                                leftMargin=0.8*inch, rightMargin=0.8*inch)
        styles = getSampleStyleSheet()
        gold = HexColor("#D4AF37")
        title_s = ParagraphStyle("T", parent=styles["Title"], fontSize=16, textColor=HexColor("#1a1a1a"),
                                 spaceAfter=4, fontName="Helvetica-Bold")
        sub_s = ParagraphStyle("S", parent=styles["Normal"], fontSize=11, textColor=HexColor("#666"), spaceAfter=12)
        body_s = ParagraphStyle("B", parent=styles["Normal"], fontSize=11, textColor=HexColor("#222"),
                                leading=16, spaceAfter=8)
        footer_s = ParagraphStyle("F", parent=styles["Normal"], fontSize=7, textColor=HexColor("#999"), alignment=1)

        story = []
        story.append(Paragraph(name, title_s))
        if company or role:
            story.append(Paragraph(f"{role} — {company}" if company and role else (role or company), sub_s))
        story.append(HRFlowable(width="100%", thickness=1, color=gold, spaceAfter=14))
        for para in body_text.split("\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(para, body_s))
        story.append(Spacer(1, 30))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#ccc"), spaceAfter=6))
        story.append(Paragraph("Generated by SaintSal Labs | saintsallabs.com", footer_s))
        doc.build(story)
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/pdf",
                                 headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'})

    elif fmt == "docx":
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.styles["Normal"].font.size = Pt(11)
        p = doc.add_heading(name, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if company or role:
            doc.add_paragraph(f"{role} — {company}" if company and role else (role or company))
        for para in body_text.split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
        p = doc.add_paragraph("\nGenerated by SaintSal Labs | saintsallabs.com")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'})

    return JSONResponse({"error": "Format must be 'pdf' or 'docx'"}, status_code=400)


# ── 10. FILE UPLOADS (Headshot / Background) ──

UPLOAD_DIR = Path("/app/backend/media_uploads/career")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

@router.post("/upload/headshot")
async def upload_headshot(file: UploadFile = File(...)):
    ext = Path(file.filename or "img.jpg").suffix.lower()
    if ext not in (".jpg", ".jpeg", ".png", ".webp"):
        return JSONResponse({"error": "Only JPG/PNG/WEBP allowed"}, status_code=400)
    fid = str(uuid.uuid4()) + ext
    dest = UPLOAD_DIR / fid
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)
    url = f"/api/career/v2/uploads/{fid}"
    return {"success": True, "url": url, "file_id": fid}

@router.post("/upload/background")
async def upload_background(file: UploadFile = File(...)):
    ext = Path(file.filename or "img.jpg").suffix.lower()
    if ext not in (".jpg", ".jpeg", ".png", ".webp"):
        return JSONResponse({"error": "Only JPG/PNG/WEBP allowed"}, status_code=400)
    fid = str(uuid.uuid4()) + ext
    dest = UPLOAD_DIR / fid
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)
    url = f"/api/career/v2/uploads/{fid}"
    return {"success": True, "url": url, "file_id": fid}

@router.get("/uploads/{file_id}")
async def serve_career_upload(file_id: str):
    fpath = UPLOAD_DIR / file_id
    if not fpath.exists():
        return JSONResponse({"error": "File not found"}, status_code=404)
    ext = fpath.suffix.lower()
    ct = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "webp": "image/webp"}.get(ext.lstrip("."), "application/octet-stream")
    return StreamingResponse(open(fpath, "rb"), media_type=ct)
